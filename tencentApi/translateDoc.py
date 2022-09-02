import json
from tencentcloud.common import credential
from tencentcloud.common.profile.client_profile import ClientProfile
from tencentcloud.common.profile.http_profile import HttpProfile
from tencentcloud.common.exception.tencent_cloud_sdk_exception import TencentCloudSDKException
from tencentcloud.tmt.v20180321 import tmt_client, models
import os
import easygui as g
from docx import Document
import time


class TencentTranslateApi:
    def __init__(self, source='auto', target='zh'):
        self.cred = credential.Credential("AKIDXu0pOhjUyB7miTzFG2Yopa7OPAgwve9O", "0lHXXHlzwc08zZsvmOlskMioeEWXJn0w")
        self.httpProfile = HttpProfile()
        self.httpProfile.endpoint = "tmt.tencentcloudapi.com"
        self.clientProfile = ClientProfile()
        self.clientProfile.httpProfile = self.httpProfile
        self.client = tmt_client.TmtClient(self.cred, "ap-shanghai", self.clientProfile)
        self.source_language = source
        self.target_language = target
        self.project_id = 0

    def multi_translate(self, text_list):
        params = {
            "Source": self.source_language,
            "Target": self.target_language,
            "ProjectId": self.project_id,
            "SourceTextList": text_list
        }
        self.project_id += 1
        req = models.TextTranslateBatchRequest()
        req.from_json_string(json.dumps(params))
        resp = self.client.TextTranslateBatch(req)
        res_json = json.loads(resp.to_json_string())
        return res_json['TargetTextList']


class TranslateWord:
    def __init__(self, file, source='auto', target='zh'):
        self.doc = Document(file)
        self.app = TencentTranslateApi(source, target)
        self.insert_translation()
        old_name = file
        rename = os.path.splitext(old_name)[0] + '转译' + '.docx'
        self.doc.save(rename)

    def insert_translation(self):
        untranslated_texts = [str(p.text) for p in self.doc.paragraphs]
        translated_texts = []
        div = 1
        while len(translated_texts) != len(untranslated_texts):
            time.sleep(0.22)
            try:
                translated_texts += self.app.multi_translate(untranslated_texts[len(translated_texts):
                                                                                len(translated_texts)
                                                                                + int((len(untranslated_texts)
                                                                                       - len(translated_texts))/div)])
                print('已翻译{}/{}\n现在分母为:{}。未翻={}'.format(len(translated_texts),
                                                        len(untranslated_texts),
                                                        div,
                                                        int((len(untranslated_texts) - len(translated_texts)))))
                div = 1
            except TencentCloudSDKException:
                print('\n现在分母为:{}'.format(div))
                div *= 2
                pass

        for p, t in zip(self.doc.paragraphs, translated_texts):
            p.text += '\n' + t


if __name__ == '__main__':
    trans = TranslateWord(g.fileopenbox('要转换的文件'), target=g.choicebox(msg="目标语言", choices=['en', 'zh']))
